#!/usr/bin/env python3
"""Batch screen resumes against a target JD and output an HR review table."""

from __future__ import annotations

import argparse
import csv
import json
import re
import sys
from pathlib import Path
from typing import Dict, Iterable, List, Optional, Tuple

try:
    from .parse_resume import extract_text, parse_resume_text
    from .recommendation_engine import (
        DEGREE_ORDER, RELEVANT_CRED_FLOOR, parse_jd, related_experiences,
        relevant_experience_months, resolve_threshold,
    )
    from .calculate_tenure import months_between, parse_date
except ImportError:
    from parse_resume import extract_text, parse_resume_text
    from recommendation_engine import (
        DEGREE_ORDER, RELEVANT_CRED_FLOOR, parse_jd, related_experiences,
        relevant_experience_months, resolve_threshold,
    )
    from calculate_tenure import months_between, parse_date


SUPPORTED_RESUME_SUFFIXES = {".pdf", ".docx", ".txt", ".md"}
IMAGE_SUFFIXES = {".jpg", ".jpeg", ".png", ".bmp", ".tiff"}
DEFAULT_WEIGHTS = {
    "relevant_experience": 0.35,
    "skill_keyword": 0.25,
    "education": 0.15,
    "stability_gap": 0.15,
    "parsing_confidence": 0.10,
}
# 推荐等级次序（仅作同分时的次级排序键；主排序按匹配分降序）。
STATUS_ORDER = {"强推荐": 0, "待审核": 1, "谨慎": 2, "不推荐": 3}


def rank_badge(n: int) -> str:
    """排序名次展示：前三名用奖牌，其余用数字。"""
    return {1: "🥇", 2: "🥈", 3: "🥉"}.get(n, str(n))


def load_text_or_path(value: str) -> str:
    try:
        path = Path(value)
        if "\n" not in value and len(value) < 512 and path.exists():
            return path.read_text(encoding="utf-8", errors="ignore")
    except OSError:
        pass
    return value


def iter_resume_files(resume_dir: Path) -> Iterable[Path]:
    for path in sorted(resume_dir.iterdir()):
        if path.is_file():
            yield path


def _dimension_for_text(text: str) -> Optional[str]:
    lowered = text.lower()
    if any(token in text for token in ["相关经验", "工作经验", "经验"]):
        return "relevant_experience"
    if any(token in text for token in ["技能", "关键词", "关键字", "命中"]):
        return "skill_keyword"
    if "学历" in text or "教育" in text:
        return "education"
    if any(token in text for token in ["稳定", "gap", "Gap", "履历"]):
        return "stability_gap"
    if any(token in text for token in ["解析", "置信", "完整度"]):
        return "parsing_confidence"
    if any(token in lowered for token in ["confidence", "parsing"]):
        return "parsing_confidence"
    return None


def _number_from_text(text: str) -> Optional[float]:
    percent = re.search(r"(\d+(?:\.\d+)?)\s*%", text)
    if percent:
        return float(percent.group(1)) / 100
    number = re.search(r"(?<!\d)(0?\.\d+|1(?:\.0+)?|\d+(?:\.\d+)?)(?!\d)", text)
    if not number:
        return None
    value = float(number.group(1))
    if value > 1:
        value = value / 100
    return value


def _normalize_weights(weights: Dict[str, float]) -> Dict[str, float]:
    merged = dict(DEFAULT_WEIGHTS)
    merged.update({key: value for key, value in weights.items() if key in DEFAULT_WEIGHTS and value >= 0})
    total = sum(merged.values())
    if total <= 0:
        return dict(DEFAULT_WEIGHTS)
    return {key: round(value / total, 4) for key, value in merged.items()}


def _weights_from_lines(lines: Iterable[str]) -> Dict[str, float]:
    weights: Dict[str, float] = {}
    for line in lines:
        text = " ".join(str(part) for part in re.split(r"[\t,，|]", line) if str(part).strip())
        dimension = _dimension_for_text(text)
        value = _number_from_text(text)
        if dimension and value is not None:
            weights[dimension] = value
    return weights


def load_weights(path: str | None) -> Tuple[Dict[str, float], List[str]]:
    warnings: List[str] = []
    if not path:
        return dict(DEFAULT_WEIGHTS), warnings

    file_path = Path(path)
    if not file_path.exists():
        warnings.append(f"权重文件不存在：{path}，已使用默认权重")
        return dict(DEFAULT_WEIGHTS), warnings

    try:
        suffix = file_path.suffix.lower()
        if suffix == ".xlsx":
            try:
                import openpyxl
            except ImportError as exc:
                raise RuntimeError("需要安装 openpyxl 才能解析 .xlsx 权重") from exc
            workbook = openpyxl.load_workbook(str(file_path), data_only=True)
            lines = []
            for sheet in workbook.worksheets:
                for row in sheet.iter_rows(values_only=True):
                    lines.append(" ".join("" if value is None else str(value) for value in row))
            return _normalize_weights(_weights_from_lines(lines)), warnings
        if suffix == ".docx":
            try:
                from docx import Document
            except ImportError as exc:
                raise RuntimeError("需要安装 python-docx 才能解析 .docx 权重") from exc
            doc = Document(str(file_path))
            lines = [paragraph.text for paragraph in doc.paragraphs]
            for table in doc.tables:
                for row in table.rows:
                    lines.append(" ".join(cell.text for cell in row.cells))
            return _normalize_weights(_weights_from_lines(lines)), warnings
        warnings.append(f"暂不支持权重文件类型：{suffix}，已使用默认权重")
    except Exception as exc:  # noqa: BLE001 - keep batch screening resilient.
        warnings.append(f"权重文件解析失败：{exc}；已使用默认权重")
    return dict(DEFAULT_WEIGHTS), warnings


def highest_degree(candidate: Dict) -> str:
    education = candidate.get("education", [])
    if not education:
        return "未解析"
    degrees = [edu.get("degree") for edu in education if edu.get("degree")]
    if not degrees:
        return "未解析"
    return max(degrees, key=lambda degree: DEGREE_ORDER.get(degree, 0))


def education_score(candidate: Dict, jd: Dict) -> float:
    required = DEGREE_ORDER.get(jd.get("min_degree") or "不限", 0)
    if required <= 0:
        return 1.0
    candidate_level = max((DEGREE_ORDER.get(edu.get("degree", ""), 0) for edu in candidate.get("education", [])), default=0)
    return 1.0 if candidate_level >= required else 0.0


def keyword_hits(candidate: Dict, jd: Dict) -> List[str]:
    jd_skills = jd.get("skills", []) or []
    text_parts = []
    for exp in candidate.get("experiences", []):
        text_parts.extend([
            exp.get("job_title", ""),
            exp.get("standardized_job_title", ""),
            exp.get("description", ""),
        ])
    text = "\n".join(part for part in text_parts if part)
    hits = [skill for skill in jd_skills if skill and skill.lower() in text.lower()]
    if not hits:
        for token in ["销售", "客户", "沟通", "成交", "回款", "外呼", "渠道", "BD", "客服"]:
            if token.lower() in text.lower():
                hits.append(token)
    return sorted(set(hits))


def stability_gap_score(candidate: Dict) -> float:
    stability = candidate.get("stability_scores", {})
    stability_score = float(stability.get("stability_score") or 0) / 100
    gap_score = float(stability.get("gap_score") or 0) / 100
    return max(0.0, min(1.0, (stability_score + gap_score) / 2))


# 注：旧 JD 加权分已下线，批量表改用单一「匹配分」(recommendation.score_100)。
# 本函数及上方权重机制（DEFAULT_WEIGHTS / load_weights / --weights）保留以便回溯，当前不参与打分。
def calculate_match_score(candidate: Dict, jd: Dict, job_title: str, weights: Dict[str, float]) -> float:
    min_months = jd.get("min_experience_months") or 6
    relevant_months = relevant_experience_months(candidate, job_title, cred_floor=RELEVANT_CRED_FLOOR)
    experience_score = min(relevant_months / max(min_months, 1), 1.0)
    skills = jd.get("skills", []) or []
    hits = keyword_hits(candidate, jd)
    if skills:
        skill_score = min(len(hits) / len(skills), 1.0)
    else:
        skill_score = min(len(hits) / 4, 1.0)
    score = (
        weights["relevant_experience"] * experience_score
        + weights["skill_keyword"] * skill_score
        + weights["education"] * education_score(candidate, jd)
        + weights["stability_gap"] * stability_gap_score(candidate)
        + weights["parsing_confidence"] * float(candidate.get("parsing_confidence") or 0)
    )
    return round(max(0.0, min(1.0, score)) * 100, 1)


def review_reasons(candidate: Dict) -> List[str]:
    reasons = []
    if candidate.get("parsing_confidence", 0) < 0.6:
        reasons.append("解析置信度低")
    for item in candidate.get("anomalies", []):
        level = item.get("level")
        if level in {"P0", "P1"} or "不参与评分" in item.get("action", ""):
            reasons.append(f"{item.get('type')}：{item.get('description')}")
    return reasons[:4]


def mismatch_items(candidate: Dict) -> List[str]:
    recommendation = candidate.get("recommendation", {})
    details = recommendation.get("details", {})
    items = list(details.get("unmet_strong_criteria", []))
    items.extend(details.get("downgrade_factors", []))
    return items[:4]


def short_comment(candidate: Dict, hits: List[str], mismatches: List[str]) -> str:
    # 简要分析：相关经验年限 + 学历 + 命中要点 + 首要风险（推荐等级/匹配分已各自成列）。
    recommendation = candidate.get("recommendation", {})
    evidence = recommendation.get("details", {}).get("evidence", {})
    parts: List[str] = []
    rel = int(evidence.get("relevant_months", 0) or 0)
    if rel > 0:
        yrs, mos = divmod(rel, 12)
        span = (f"{yrs}年" if yrs else "") + (f"{mos}个月" if mos and not yrs else "")
        parts.append(f"{span}相关经验")
    degree = highest_degree(candidate)
    if degree and degree not in ("未解析", "未说明", "-"):
        parts.append(degree)
    parts.append("命中" + "、".join(hits[:3]) if hits else "命中点较少")
    if mismatches:
        parts.append("风险：" + mismatches[0])
    base = "；".join(parts)[:70]
    # 有 P0 时把具体内容加粗拼进摘要，确保只看摘要列也能看到（在截断后追加，保证 ** 配对）。
    p0_remark = recommendation.get("p0_remark", "")
    p0_text = f"；**⚠️P0：{p0_remark[:40]}**" if p0_remark else ""
    return base + p0_text


def row_from_candidate(path: Path, candidate: Dict, jd: Dict, job_title: str, weights: Dict[str, float]) -> Dict:
    recommendation = candidate.get("recommendation", {})
    # 解析置信度已并入匹配分（D4 维度），不再单独硬降级；等级由匹配分单调推导。
    status = recommendation.get("status", "不推荐")
    tier = recommendation.get("tier", status)
    hits = keyword_hits(candidate, jd)
    mismatches = mismatch_items(candidate)
    reasons = review_reasons(candidate)
    row = {
        "source_file": path.name,
        "candidate": candidate.get("basic_info", {}).get("name") or path.stem,
        "evidence_score": recommendation.get("score_100", 0),
        "tier": tier,
        "tier_display": recommendation.get("tier_display", tier),
        "status": status,
        "recommendation_score": recommendation.get("score", 0),
        "parsing_confidence": candidate.get("parsing_confidence", 0),
        "education": highest_degree(candidate),
        "relevant_experience_months": relevant_experience_months(candidate, job_title, cred_floor=RELEVANT_CRED_FLOOR),
        "keyword_hits": "、".join(hits) if hits else "-",
        "mismatch_items": "；".join(mismatches) if mismatches else "-",
        "p0_remark": recommendation.get("p0_remark", ""),
        "review_reasons": "；".join(reasons) if reasons else "-",
        "comment": short_comment(candidate, hits, mismatches),
        "candidate_card": candidate,
    }
    return row


def unparsed_row(path: Path, reason: str) -> Dict:
    return {
        "source_file": path.name,
        "candidate": path.stem,
        "evidence_score": 0,
        "tier": "不推荐",
        "tier_display": "🔴 不推荐（未解析）",
        "status": "不推荐",
        "recommendation_score": 0,
        "parsing_confidence": 0,
        "education": "未解析",
        "relevant_experience_months": 0,
        "keyword_hits": "-",
        "mismatch_items": "-",
        "p0_remark": "",
        "review_reasons": reason,
        "comment": "文件未完成解析，需人工处理",
        "candidate_card": None,
    }


def sort_rows(rows: List[Dict]) -> List[Dict]:
    # 主排序：匹配分降序（等级已随匹配分单调，脱钩消失）；同分再按等级、旧内部分。
    ordered = sorted(
        rows,
        key=lambda row: (
            -float(row.get("evidence_score") or 0),
            STATUS_ORDER.get(row.get("tier") or row.get("status"), 9),
            -float(row.get("recommendation_score") or 0),
        ),
    )
    for idx, row in enumerate(ordered, 1):
        row["rank"] = idx
    return ordered


# Markdown 排序表（对齐截图版式，精简自解释）：排序 | 候选人 | 匹配分 | 推荐等级 | 学历 | 相关经验 | 简要分析
TABLE_FIELDS = [
    "rank",
    "candidate",
    "evidence_score",
    "tier_display",
    "education",
    "relevant_experience_months",
    "comment",
]
# CSV 导出保留更全的明细列，供需要深挖的 HR 使用。
CSV_FIELDS = [
    "rank",
    "candidate",
    "evidence_score",
    "tier",
    "parsing_confidence",
    "education",
    "relevant_experience_months",
    "keyword_hits",
    "mismatch_items",
    "p0_remark",
    "review_reasons",
    "comment",
]
FIELD_LABELS = {
    "rank": "排序",
    "candidate": "候选人",
    "evidence_score": "匹配分",
    "tier": "推荐等级",
    "tier_display": "推荐等级",
    "parsing_confidence": "解析置信度",
    "education": "学历",
    "relevant_experience_months": "相关经验(月)",
    "keyword_hits": "命中关键词",
    "mismatch_items": "不匹配项",
    "p0_remark": "P0高亮",
    "review_reasons": "人工复核原因",
    "comment": "简要分析",
}


def _cell(row: Dict, field: str) -> str:
    if field == "rank":
        return rank_badge(int(row.get("rank") or 0))
    return str(row.get(field, "-")).replace("\n", " ").replace("|", "/")


def markdown_table(rows: List[Dict], warnings: List[str]) -> str:
    lines: List[str] = []
    if warnings:
        lines.append("\n".join(f"> 警告：{warning}" for warning in warnings))
        lines.append("")
    header = [FIELD_LABELS[field] for field in TABLE_FIELDS]
    lines.append("| " + " | ".join(header) + " |")
    lines.append("| " + " | ".join("---" for _ in header) + " |")
    for row in rows:
        values = [_cell(row, field) for field in TABLE_FIELDS]
        lines.append("| " + " | ".join(values) + " |")
    return "\n".join(lines)


def write_output(rows: List[Dict], output: str | None, warnings: List[str]) -> None:
    if not output:
        print(markdown_table(rows, warnings))
        return

    path = Path(output)
    suffix = path.suffix.lower()
    if suffix == ".json":
        payload = {"warnings": warnings, "rows": rows}
        path.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")
        return
    if suffix == ".csv":
        with path.open("w", newline="", encoding="utf-8-sig") as csvfile:
            writer = csv.DictWriter(csvfile, fieldnames=CSV_FIELDS, extrasaction="ignore")
            writer.writerow({field: FIELD_LABELS[field] for field in CSV_FIELDS})
            for row in rows:
                writer.writerow({field: row.get(field, "") for field in CSV_FIELDS})
        return
    path.write_text(markdown_table(rows, warnings), encoding="utf-8")


def run_batch(resume_dir: Path, jd_text: str, job_title: str, weights: Dict[str, float], pass_threshold: Optional[float] = None) -> List[Dict]:
    jd = parse_jd(jd_text)
    effective_job_title = job_title or jd.get("job_title") or "销售"
    rows: List[Dict] = []
    files = list(iter_resume_files(resume_dir))
    if not files:
        raise RuntimeError(f"简历目录为空：{resume_dir}")

    for path in files:
        suffix = path.suffix.lower()
        if suffix in IMAGE_SUFFIXES:
            # 先尝试 OCR；依赖未装/失败时优雅降级为人工处理（不中断批量）。
            try:
                text = extract_text(str(path))
                candidate = parse_resume_text(text, jd_text=jd_text, job_title=effective_job_title, pass_threshold=pass_threshold)
                row = row_from_candidate(path, candidate, jd, effective_job_title, weights)
                note = "图片经 OCR 解析，建议人工二次确认"
                existing = row.get("review_reasons", "")
                row["review_reasons"] = note if existing in ("", "-") else f"{existing}；{note}"
                rows.append(row)
            except Exception as exc:  # noqa: BLE001 - OCR 不可用/失败则退回人工。
                rows.append(unparsed_row(path, f"图片简历需人工处理：{exc}"))
            continue
        if suffix not in SUPPORTED_RESUME_SUFFIXES:
            rows.append(unparsed_row(path, f"暂不支持文件类型：{suffix}"))
            continue
        try:
            text = extract_text(str(path))
            candidate = parse_resume_text(text, jd_text=jd_text, job_title=effective_job_title, pass_threshold=pass_threshold)
            rows.append(row_from_candidate(path, candidate, jd, effective_job_title, weights))
        except Exception as exc:  # noqa: BLE001 - one bad resume must not stop the batch.
            rows.append(unparsed_row(path, f"解析失败：{exc}"))
    return sort_rows(rows)


def main() -> None:
    # Windows 中文控制台(cp936)编不出表格里的 ✅ 等符号会让 print 崩溃；强制 utf-8、容错替换。
    for stream in (sys.stdout, sys.stderr):
        try:
            stream.reconfigure(encoding="utf-8", errors="replace")
        except Exception:  # noqa: BLE001
            pass
    parser = argparse.ArgumentParser(description="Batch screen resumes against a target JD.")
    parser.add_argument("resume_dir", help="Directory containing resumes")
    parser.add_argument("--jd", required=True, help="JD text or path to a JD text file")
    parser.add_argument("--job-title", default="", help="Target job title")
    parser.add_argument("--weights", default="", help="（已弃用，保留兼容）旧 JD 匹配分权重文件；当前评分改为单一证据强度分，不再生效")
    parser.add_argument(
        "--pass-threshold",
        type=float,
        default=None,
        help="强推荐门槛（匹配分 0-100），默认 75；待审核=门槛-10、谨慎=门槛-20。可用环境变量 HR_PASS_THRESHOLD 覆盖",
    )
    parser.add_argument("--output", default="", help="Optional .md/.csv/.json output path")
    args = parser.parse_args()

    resume_dir = Path(args.resume_dir)
    if not resume_dir.exists() or not resume_dir.is_dir():
        raise SystemExit(f"简历目录不存在或不是目录：{resume_dir}")
    jd_text = load_text_or_path(args.jd)
    if not jd_text.strip():
        raise SystemExit("JD 不能为空：请通过 --jd 提供 JD 文件路径或文本")
    weights, warnings = load_weights(args.weights)
    if args.weights:
        print("提示：评分已改为单一证据强度分，--weights 不再参与打分（仅保留兼容）。", file=sys.stderr)
    for warning in warnings:
        print(f"警告：{warning}", file=sys.stderr)
    threshold = resolve_threshold(args.pass_threshold)
    try:
        rows = run_batch(resume_dir, jd_text, args.job_title, weights, pass_threshold=threshold)
    except RuntimeError as exc:
        raise SystemExit(str(exc)) from exc
    write_output(rows, args.output or None, warnings)


if __name__ == "__main__":
    main()
